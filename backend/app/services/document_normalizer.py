from __future__ import annotations

import io
from dataclasses import dataclass
from pathlib import Path

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas

try:  # pragma: no cover - optional import failures handled at runtime
    from PIL import Image
except ImportError:  # pragma: no cover
    Image = None  # type: ignore[misc, assignment]

try:  # pragma: no cover
    from docx import Document as DocxDocument
except ImportError:  # pragma: no cover
    DocxDocument = None  # type: ignore[misc, assignment]


IMAGE_CONTENT_TYPES = {
    "image/png",
    "image/jpeg",
    "image/jpg",
    "image/gif",
    "image/webp",
    "image/tiff",
}

IMAGE_EXTENSIONS = {
    ".png",
    ".jpg",
    ".jpeg",
    ".gif",
    ".webp",
    ".tiff",
    ".tif",
}


@dataclass
class NormalizedDocument:
    pdf_bytes: bytes
    filename: str
    converted: bool


def normalize_to_pdf(filename: str | None, content_type: str | None, data: bytes) -> NormalizedDocument:
    """Normalize uploaded content into a PDF representation."""

    safe_name = _sanitize_filename(filename or "document")
    lower_ext = Path(safe_name).suffix.lower()
    content_type = (content_type or "").lower()

    if lower_ext == ".pdf" or content_type == "application/pdf" or data.startswith(b"%PDF"):
        pdf_name = _ensure_pdf_extension(safe_name)
        return NormalizedDocument(pdf_bytes=data, filename=pdf_name, converted=False)

    if lower_ext == ".docx" or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        pdf_bytes = _convert_docx_to_pdf(data)
        pdf_name = _ensure_pdf_extension(safe_name)
        return NormalizedDocument(pdf_bytes=pdf_bytes, filename=pdf_name, converted=True)

    if lower_ext in IMAGE_EXTENSIONS or content_type in IMAGE_CONTENT_TYPES:
        pdf_bytes = _convert_image_to_pdf(data)
        pdf_name = _ensure_pdf_extension(safe_name)
        return NormalizedDocument(pdf_bytes=pdf_bytes, filename=pdf_name, converted=True)

    raise ValueError("Unsupported file type for automatic normalization")


def _sanitize_filename(name: str) -> str:
    stem = Path(name).stem or "document"
    return stem.replace(" ", "_")


def _ensure_pdf_extension(name: str) -> str:
    stem = Path(name).stem or "document"
    return f"{stem}.pdf"


def _convert_docx_to_pdf(data: bytes) -> bytes:
    if DocxDocument is None:  # pragma: no cover - defensive guard
        raise ValueError("DOCX conversion unavailable: python-docx not installed")

    buffer = io.BytesIO(data)
    try:
        document = DocxDocument(buffer)
    except Exception as exc:  # pragma: no cover - docx parsing failure
        raise ValueError("Invalid DOCX file") from exc

    pdf_buffer = io.BytesIO()
    pdf_canvas = canvas.Canvas(pdf_buffer, pagesize=LETTER)
    width, height = LETTER

    margin = 54  # 0.75 inch
    y_cursor = height - margin

    def write_line(line: str) -> None:
        nonlocal y_cursor
        if y_cursor <= margin:
            pdf_canvas.showPage()
            y_cursor = height - margin
        pdf_canvas.drawString(margin, y_cursor, line)
        y_cursor -= 14

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            y_cursor -= 7
            continue
        for line in _wrap_text(text, max_chars=90):
            write_line(line)
        y_cursor -= 7

    pdf_canvas.save()
    return pdf_buffer.getvalue()


def _convert_image_to_pdf(data: bytes) -> bytes:
    if Image is None:  # pragma: no cover - defensive guard
        raise ValueError("Image conversion unavailable: Pillow not installed")

    with Image.open(io.BytesIO(data)) as img:
        if img.mode in {"RGBA", "P"}:
            img = img.convert("RGB")
        dpi = img.info.get("dpi", (72, 72))
        dpi_x = dpi[0] or 72
        dpi_y = dpi[1] or 72
        width_pts = img.width / dpi_x * 72
        height_pts = img.height / dpi_y * 72

        pdf_buffer = io.BytesIO()
        pdf_canvas = canvas.Canvas(pdf_buffer, pagesize=(width_pts, height_pts))
        pdf_canvas.drawImage(ImageReader(img), 0, 0, width=width_pts, height=height_pts, preserveAspectRatio=True)
        pdf_canvas.save()
        return pdf_buffer.getvalue()


def _wrap_text(text: str, max_chars: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current: list[str] = []
    length = 0

    for word in words:
        prospective = length + len(word) + (1 if current else 0)
        if prospective > max_chars and current:
            lines.append(" ".join(current))
            current = [word]
            length = len(word)
        else:
            current.append(word)
            length = prospective

    if current:
        lines.append(" ".join(current))
    return lines
